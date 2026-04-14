import streamlit as st
import pandas as pd
import plotly.express as px
import os
import warnings
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

# --- 1. CONFIGURACIÓN ---
warnings.filterwarnings("ignore")
st.set_page_config(page_title="Fundación del Valle", layout="wide")

AZUL_FDV = "#003366"
NARANJA_FDV = "#E67E22"
VERDE_ACTIVO = "#28a745"
AMARILLO_MAPA = "#FFB300"

# Logo y CSS Personalizado
LOGO_URL = "https://coordinadoraongd.org/wp-content/uploads/2016/04/fundacion_del_valle.jpg"

st.markdown(f"""
    <style>
    .stApp {{ background-color: #FFFFFF; }}
    /* Logo arriba a la derecha */
    .logo-container {{
        position: absolute;
        top: -50px;
        right: 0px;
        z-index: 1000;
    }}
    .logo-img {{
        width: 180px;
    }}
    h1, h2, h3 {{ color: {AZUL_FDV} !important; font-family: 'Arial'; }}
    .stMetric {{ background-color: #F8F9FA; padding: 15px; border-radius: 10px; border-left: 5px solid {NARANJA_FDV}; }}
    .ficha-box {{ border: 1px solid #DDD; padding: 20px; border-radius: 10px; margin-bottom: 20px; background-color: #FAFAFA; }}
    </style>
    <div class="logo-container">
        <img src="{LOGO_URL}" class="logo-img">
    </div>
    """, unsafe_allow_html=True)

def fmt_euro(valor):
    try:
        # Formato: 1.234.567 €
        return f"{int(round(float(valor))):,.0f}".replace(",", ".") + " €"
    except:
        return "0 €"

def fmt_numero(valor):
    try:
        return f"{int(valor):,.0f}".replace(",", ".")
    except:
        return "0"

# --- 2. CARGA DE DATOS ---
@st.cache_data
def load_data():
    folder = os.path.dirname(os.path.abspath(__file__))
    target = os.path.join(folder, "datos.csv.xlsx")
    if not os.path.exists(target): return pd.DataFrame()
    try:
        df = pd.read_csv(target, sep=",", skiprows=1, encoding='utf-8')
    except:
        df = pd.read_excel(target, skiprows=1)
    
    df.columns = [str(c).strip() for c in df.columns]
    
    cols_txt = ["Sector 1", "PAIS", "SOCIO LOCAL/CONTRAPARTE 1", "Título", "Financiador"]
    for c in cols_txt:
        if c in df.columns:
            df[c] = df[c].fillna('').astype(str).replace(['nan', 'None', 'N/A'], '').str.strip()

    num_cols = ["SUBVENCIÓN", "COSTE TOTAL", "B. Directos (Nº)", "B. Indirectos (Nº)"]
    for n in num_cols:
        if n in df.columns:
            df[n] = df[n].astype(str).str.replace(r'[^\d,.-]', '', regex=True).str.replace(',', '.')
            df[n] = pd.to_numeric(df[n], errors='coerce').fillna(0)

    df["Año_Num"] = pd.to_numeric(df["AÑO"], errors='coerce').fillna(0).astype(int)
    
    if 'FIN' in df.columns:
        df['FECHA_FIN_DT'] = pd.to_datetime(df['FIN'], errors='coerce')
        hoy = pd.to_datetime(datetime.now().date())
        status_paises = df.groupby("PAIS")["FECHA_FIN_DT"].max().reset_index()
        status_paises["Estado"] = status_paises["FECHA_FIN_DT"].apply(
            lambda x: "Activo" if pd.notnull(x) and x >= hoy else "Histórico"
        )
        df = df.merge(status_paises[["PAIS", "Estado"]], on="PAIS", how="left")
    return df

df = load_data()

if not df.empty:
    # --- 3. BARRA LATERAL (FILTROS) ---
    with st.sidebar:
        st.header("🔍 Filtros Globales")
        st.markdown("---")
        f_pais = st.multiselect("📍 País", sorted([str(x) for x in df["PAIS"].unique() if str(x) != '']))
        f_año = st.multiselect("📅 Año", sorted([int(x) for x in df["Año_Num"].unique() if x > 0], reverse=True))
        f_sector = st.multiselect("🎯 Sector", sorted([str(x) for x in df["Sector 1"].unique() if str(x) != '']))
        f_socio = st.multiselect("🤝 Socio Local", sorted([str(x) for x in df["SOCIO LOCAL/CONTRAPARTE 1"].unique() if str(x) != '']))
        f_finan = st.multiselect("💰 Financiador", sorted([str(x) for x in df["Financiador"].unique() if str(x) != '']))

    df_f = df.copy()
    if f_pais: df_f = df_f[df_f["PAIS"].isin(f_pais)]
    if f_año: df_f = df_f[df_f["Año_Num"].isin(f_año)]
    if f_sector: df_f = df_f[df_f["Sector 1"].isin(f_sector)]
    if f_socio: df_f = df_f[df_f["SOCIO LOCAL/CONTRAPARTE 1"].isin(f_socio)]
    if f_finan: df_f = df_f[df_f["Financiador"].isin(f_finan)]

    # --- 4. DASHBOARD ---
    st.title("🌍 Impacto Global - Fundación del Valle")
    
    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("Proyectos", f"{len(df_f)}")
    m2.metric("Subvención", fmt_euro(df_f['SUBVENCIÓN'].sum()))
    m3.metric("Coste Total", fmt_euro(df_f['COSTE TOTAL'].sum()))
    m4.metric("Ben. Directos", fmt_numero(df_f['B. Directos (Nº)'].sum()))
    m5.metric("Ben. Indirectos", fmt_numero(df_f['B. Indirectos (Nº)'].sum()))

    st.divider()
    
    # --- MAPA CON HOVER PERSONALIZADO ---
    st.subheader("📍 Presencia Institucional Global")
    df_mapa = df_f.groupby("PAIS").agg({
        "SUBVENCIÓN": "sum", 
        "B. Directos (Nº)": "sum", 
        "Estado": "first"
    }).reset_index()

    # Preparar etiquetas para el mapa
    df_mapa["Importe_Etiqueta"] = df_mapa["SUBVENCIÓN"].apply(fmt_euro)
    df_mapa["Ben_Etiqueta"] = df_mapa["B. Directos (Nº)"].apply(fmt_numero)

    fig_map = px.choropleth(
        df_mapa, 
        locations="PAIS", 
        locationmode='country names', 
        color="Estado",
        color_discrete_map={"Activo": VERDE_ACTIVO, "Histórico": AMARILLO_MAPA}, 
        projection="natural earth"
    )

    fig_dots = px.scatter_geo(
        df_mapa, 
        locations="PAIS", 
        locationmode='country names', 
        size="SUBVENCIÓN",
        hover_name="PAIS", 
        custom_data=["Importe_Etiqueta", "Ben_Etiqueta"], 
        size_max=30
    )

    # Configuración del cursor (Hover)
    fig_dots.update_traces(
        hovertemplate="<b>%{hovertext}</b><br>" +
                      "Importe: %{customdata[0]}<br>" +
                      "Total Beneficiarios Directos: %{customdata[1]}<extra></extra>",
        marker=dict(color=AZUL_FDV, opacity=0.7)
    )

    fig_map.add_trace(fig_dots.data[0])
    fig_map.update_geos(showcountries=True, countrycolor="white")
    fig_map.update_layout(margin={"r":0,"t":0,"l":0,"b":0}, height=600, showlegend=False)
    st.plotly_chart(fig_map, use_container_width=True)

    st.divider()

    # Gráficas de Sectores
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("👥 Beneficiarios por Sector")
        df_b = df_f.groupby("Sector 1")["B. Directos (Nº)"].sum().reset_index().sort_values("B. Directos (Nº)", ascending=True)
        fig_bar = px.bar(df_b, x="B. Directos (Nº)", y="Sector 1", orientation='h', 
                         color="Sector 1", color_discrete_sequence=px.colors.qualitative.Pastel)
        fig_bar.update_layout(showlegend=False, height=400)
        st.plotly_chart(fig_bar, use_container_width=True)
    with c2:
        st.subheader("🍩 Inversión por Sector")
        df_p = df_f.groupby("Sector 1")["COSTE TOTAL"].sum().reset_index()
        fig_pie = px.pie(df_p, values="COSTE TOTAL", names="Sector 1", hole=0.5,
                         color_discrete_sequence=px.colors.qualitative.Pastel)
        fig_pie.update_layout(height=400)
        st.plotly_chart(fig_pie, use_container_width=True)

    st.divider()
    
    # --- GENERADOR DE FICHAS ---
    st.subheader("📋 Generador de Fichas de Proyecto")
    proyectos_opciones = sorted([str(x) for x in df_f["Título"].unique() if str(x) != ''])
    seleccionados = st.multiselect("Selecciona proyectos para visualizar y descargar:", proyectos_opciones)

    if seleccionados:
        def crear_word_pro(lista_titulos, dataframe):
            doc = Document()
            for idx, titulo in enumerate(lista_titulos):
                p = dataframe[dataframe["Título"] == titulo].iloc[0]
                t = doc.add_heading(str(p['Título']).upper(), 0)
                t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                table = doc.add_table(rows=3, cols=2)
                table.style = 'Table Grid'
                def fill_cell(r, c, label, value):
                    cell = table.cell(r, c)
                    para = cell.paragraphs[0]
                    run_label = para.add_run(f"{label}: ")
                    run_label.bold = True
                    para.add_run(str(value))

                fill_cell(0, 0, "PAÍS", p['PAIS'])
                fill_cell(0, 1, "AÑO", int(p['Año_Num']))
                fill_cell(1, 0, "SOCIO", p['SOCIO LOCAL/CONTRAPARTE 1'])
                fill_cell(1, 1, "FINANCIACIÓN", fmt_euro(p['SUBVENCIÓN']))
                fill_cell(2, 0, "BENEFICIARIOS", f"{fmt_numero(p['B. Directos (Nº)'])} Directos")
                fill_cell(2, 1, "FINANCIADOR", p['Financiador'])

                doc.add_paragraph("\n")
                doc.add_heading("OBJETIVO GENERAL (OG)", level=1)
                doc.add_paragraph(str(p.get('OBJETIVO GENERAL (OG)', 'N/A')))
                doc.add_heading("OBJETIVO ESPECÍFICO (OE)", level=1)
                doc.add_paragraph(str(p.get('OBJETIVO ESPECÍFICO (OE)', 'N/A')))

                if idx < len(lista_titulos) - 1:
                    doc.add_page_break()
            
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        st.download_button(
            label=f"📥 Descargar {len(seleccionados)} proyectos en Word",
            data=crear_word_pro(seleccionados, df_f),
            file_name=f"Fichas_Tecnicas_FDV.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        for proy in seleccionados:
            p = df_f[df_f["Título"] == proy].iloc[0]
            st.markdown(f"""<div class="ficha-box">
                <h3>{p['Título']}</h3>
                <p><b>📍 País:</b> {p['PAIS']} | <b>📅 Año:</b> {int(p['Año_Num'])}</p>
                <p><b>🤝 Socio:</b> {p['SOCIO LOCAL/CONTRAPARTE 1']} | <b>💰 Financiador:</b> {p['Financiador']}</p>
                <p><b>💶 Subvención:</b> {fmt_euro(p['SUBVENCIÓN'])} | <b>👥 Impacto:</b> {fmt_numero(p['B. Directos (Nº)'])} Directos</p>
                <hr>
                <p><b>Objetivo General (OG):</b><br>{p.get('OBJETIVO GENERAL (OG)', 'N/A')}</p>
                <p><b>Objetivo Específico (OE):</b><br>{p.get('OBJETIVO ESPECÍFICO (OE)', 'N/A')}</p>
            </div>""", unsafe_allow_html=True)

else:
    st.error("Archivo de datos no encontrado.")
